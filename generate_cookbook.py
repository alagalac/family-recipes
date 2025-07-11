from fpdf import FPDF
from docx import Document
import yaml
import os

# Load the cookbook structure
def load_structure(filename):
    with open(filename, 'r') as f:
        return yaml.safe_load(f)

# Load a single recipe
def load_recipe(filename):
    with open(filename, 'r') as f:
        return yaml.safe_load(f)

# Convert recipe YAML to markdown format
def recipe_to_markdown(recipe):
    markdown_content = f"{recipe['title']}\n\n"
    markdown_content += f"**Prep Time**: {recipe['prep_time']}  |  **Cook Time**: {recipe['cook_time']}  |  **Servings**: {recipe['servings']}\n\n"
    
    markdown_content += "## Ingredients\n"
    for ingredient in recipe['ingredients']:
        markdown_content += f"- {ingredient}\n"

    markdown_content += "\n## Instructions\n"
    for i, step in enumerate(recipe['instructions'], start=1):
        markdown_content += f"{i}. {step}\n"
    
    if recipe['notes']:
        markdown_content += f"\n## Notes\n{recipe['notes']}\n"

    return markdown_content

# Generate the fancy cookbook PDF with custom layout
def generate_pdf(cookbook_structure, recipes_folder):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Set font for recipe name (Title)
    pdf.set_font('Arial', 'B', 36)
    pdf.cell(200, 14, 'My Cookbook', 0, 1, 'C')  # Cookbook Title
    pdf.ln(10)
    pdf.add_page()

    for section in cookbook_structure['sections']:
        # Add Section Title
        pdf.set_font('Arial', 'B', 26)
        pdf.set_text_color(30, 60, 120)
        pdf.cell(200, 12, section['name'], ln=True, align='L')
        pdf.set_text_color(0, 0, 0)
        pdf.ln(5)
        pdf.add_page()

        pdf.set_font('Arial', '', 12)

        # Check if the section has recipes
        if not section.get('recipes'):
            continue

        for recipe_name in section['recipes']:
            recipe_file = os.path.join(recipes_folder, f"{recipe_name}.yaml")
            recipe = load_recipe(recipe_file)

            # Recipe Title (Big)
            pdf.set_font('Times', 'B', 26)
            pdf.set_text_color(80, 30, 30)
            pdf.cell(200, 12, recipe['title'], ln=True, align='L')
            pdf.set_text_color(0, 0, 0)
            pdf.ln(5)

            # Add Commentary section if present (directly under title)
            if recipe.get('commentary'):
                pdf.set_font('Times', 'I', 13)
                pdf.set_text_color(100, 80, 40)
                pdf.multi_cell(0, 10, recipe['commentary'])
                pdf.set_text_color(0, 0, 0)
                pdf.set_font('Helvetica', '', 12)
                pdf.ln(3)

            # Horizontal row for prep time, cook time, and servings
            pdf.set_font('Helvetica', '', 12)
            pdf.cell(60, 10, f"Prep Time: {recipe['prep_time']}", align='L')
            pdf.cell(60, 10, f"Cook Time: {recipe['cook_time']}", align='C')
            pdf.cell(60, 10, f"Servings: {recipe['servings']}", align='R')
            pdf.ln(12)

            # Split the page into two columns: ingredients and instructions
            pdf.set_font('Helvetica', 'B', 13)
            
            # Left column for Ingredients
            x_left = pdf.get_x()
            y_top = pdf.get_y()
            pdf.set_fill_color(230, 240, 255)
            pdf.cell(80, 10, 'Ingredients', border=1, align='C', fill=True)
            pdf.ln(10)
            pdf.set_font('Arial', '', 11)  # was 'Courier', '', 11
            pdf.set_fill_color(255, 255, 255)
            for ingredient in recipe['ingredients']:
                pdf.multi_cell(80, 8, ingredient, border=0, fill=True)
            y_after_ingredients = pdf.get_y()

            # Right column for Instructions
            pdf.set_xy(x_left + 90, y_top)
            pdf.set_font('Helvetica', 'B', 13)
            pdf.set_fill_color(230, 255, 230)
            pdf.cell(110, 10, 'Instructions', border=1, align='C', fill=True)
            pdf.ln(10)
            pdf.set_font('Arial', '', 11)  # was 'Courier', '', 11
            pdf.set_fill_color(255, 255, 255)
            for i, step in enumerate(recipe['instructions'], start=1):
                pdf.set_x(x_left + 90)
                pdf.multi_cell(110, 8, f"{i}. {step}", border=0, fill=True)

            # Draw a subtle vertical line between columns
            pdf.set_draw_color(180, 180, 180)
            pdf.line(x_left + 85, y_top, x_left + 85, max(y_after_ingredients, pdf.get_y()))

            # Move Y to the lower of the two columns for next content
            pdf.set_y(max(y_after_ingredients, pdf.get_y()))

            # Add Notes section if present (handle list or string)
            notes = recipe.get('notes')
            if notes:
                pdf.ln(5)
                pdf.set_font('Helvetica', 'B', 12)
                pdf.set_text_color(60, 60, 120)
                pdf.cell(0, 10, 'Notes', ln=True)
                pdf.set_font('Helvetica', 'I', 12)
                pdf.set_text_color(40, 40, 40)
                pdf.set_fill_color(245, 245, 220)
                if isinstance(notes, list):
                    for note in notes:
                        pdf.multi_cell(0, 10, f"- {note}", border=0, fill=True)
                else:
                    pdf.multi_cell(0, 10, notes, border=1, fill=True)
                pdf.set_text_color(0, 0, 0)
                pdf.set_font('Helvetica', '', 12)
                pdf.set_fill_color(255, 255, 255)

            # Add Attribution section if present
            if recipe.get('attribution'):
                pdf.ln(2)
                pdf.set_font('Helvetica', 'I', 10)
                pdf.set_text_color(120, 120, 120)
                pdf.cell(0, 8, f"Attribution: {recipe['attribution']}", ln=True)
                pdf.set_text_color(0, 0, 0)
                pdf.set_font('Helvetica', '', 12)

            # Page break after each recipe
            pdf.add_page()

    pdf.output("cookbook_fancy_layout.pdf")


# Generate the fancy Word document with custom layout
def generate_word(cookbook_structure, recipes_folder):
    doc = Document()

    # Add a title to the document
    doc.add_heading('My Fancy Cookbook', 0)

    for section in cookbook_structure['sections']:
        # Add Section Title
        doc.add_heading(section['name'], level=1)

        # Check if the section has recipes
        if not section.get('recipes'):
            continue

        for recipe_name in section['recipes']:
            recipe_file = os.path.join(recipes_folder, f"{recipe_name}.yaml")
            recipe = load_recipe(recipe_file)

            # Recipe Title (Big)
            doc.add_heading(recipe['title'], level=2)

            # Add Commentary section if present (directly under title)
            if recipe.get('commentary'):
                doc.add_paragraph(recipe['commentary'], style='Intense Quote')

            # Prep Time, Cook Time, Servings in one row (horizontal)
            doc.add_paragraph(f"Prep Time: {recipe['prep_time']}  |  Cook Time: {recipe['cook_time']}  |  Servings: {recipe['servings']}")

            # Create a table for Ingredients and Instructions (2 columns)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            cells = table.rows[0].cells

            # Left column for Ingredients
            cells[0].text = 'Ingredients'
            for ingredient in recipe['ingredients']:
                cells[0].add_paragraph(ingredient)

            # Right column for Instructions
            cells[1].text = 'Instructions'
            for i, step in enumerate(recipe['instructions'], start=1):
                cells[1].add_paragraph(f"{i}. {step}")

            # Add Notes section if present (handle list or string)
            notes = recipe.get('notes')
            if notes:
                doc.add_paragraph('Notes', style='Heading 3')
                if isinstance(notes, list):
                    for note in notes:
                        doc.add_paragraph(note, style='List Bullet')
                else:
                    doc.add_paragraph(notes)

            # Add Attribution section if present
            if recipe.get('attribution'):
                doc.add_paragraph(f"Attribution: {recipe['attribution']}", style='Intense Quote')

            # Add a page break after each recipe
            doc.add_page()

    doc.save('fancy_cookbook_layout.docx')


# Main script to choose the format and generate the cookbook
if __name__ == "__main__":
    # Load the cookbook structure
    structure_filename = 'cookbook_structure.yaml'
    cookbook_structure = load_structure(structure_filename)
    recipes_folder = './recipes'

    # Choose output format: 'pdf' or 'word'
    output_format = 'pdf'  # Change to 'word' to generate a Word document

    if output_format == 'pdf':
        generate_pdf(cookbook_structure, recipes_folder)
    elif output_format == 'word':
        generate_word(cookbook_structure, recipes_folder)
    else:
        print("Unsupported format. Choose 'pdf' or 'word'.")
